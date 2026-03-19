import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
import re

def get_law_paragraph(gesetz, num):
    url = f"https://www.gesetze-im-internet.de/{gesetz}/__{num}.html"
    try:
        # Wir nutzen einen User-Agent, damit der Server uns nicht sofort blockt
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            title = soup.find('span', class_='jnnormtitel').get_text(strip=True)
            paragraphs = soup.find_all('div', class_='jurAbsatz')
            text_content = "\n".join([p.get_text(strip=True) for p in paragraphs])
            return title, text_content
        return None, f"Paragraph {num} im {gesetz.upper()} nicht gefunden."
    except Exception as e:
        return None, f"Fehler: {str(e)}"

# --- Streamlit Oberfläche ---
st.set_page_config(page_title="Gesetze-Generator", page_icon="⚖️")
st.title("⚖️ Gesetze-Generator (BGB & HGB)")

gesetz_wahl = st.selectbox("Wähle das Gesetz:", ["BGB", "HGB"])
user_input = st.text_input("Paragrafen-Nummern (z.B. 433, 823):")

if st.button("Inhalte abrufen & Word-Datei erstellen"):
    if user_input:
        nums = [n.strip() for n in re.split(r'[ ,;]+', user_input) if n.strip()]
        doc = Document()
        doc.add_heading(f'{gesetz_wahl} Auszüge', 0)
        
        found_count = 0
        for num in nums:
            title, content = get_law_paragraph(gesetz_wahl.lower(), num)
            if title:
                doc.add_heading(title, level=1)
                doc.add_paragraph(content)
                st.success(f"Gefunden: {title}")
                found_count += 1
            else:
                st.error(content)

        if found_count > 0:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="📥 Word-Datei herunterladen",
                data=buffer,
                file_name=f"{gesetz_wahl}_Export.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Bitte gib Nummern ein.")